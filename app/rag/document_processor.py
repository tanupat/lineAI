import os
from typing import List
from pathlib import Path
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from app.core.config import get_settings


class DocumentProcessor:
    """Process and chunk documents for RAG."""

    SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx", ".md", ".json", ".csv"}

    def __init__(self):
        settings = get_settings()
        self.chunk_size = settings.chunk_size
        self.chunk_overlap = settings.chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )

    def is_supported(self, filename: str) -> bool:
        """Check if the file type is supported."""
        ext = Path(filename).suffix.lower()
        return ext in self.SUPPORTED_EXTENSIONS

    async def process_file(self, file_path: str, filename: str) -> List[Document]:
        """
        Process a file and return chunked documents.

        Args:
            file_path: Path to the file
            filename: Original filename

        Returns:
            List of Document objects with chunks
        """
        ext = Path(filename).suffix.lower()

        if ext == ".txt" or ext == ".md":
            content = await self._read_text_file(file_path)
        elif ext == ".pdf":
            content = await self._read_pdf_file(file_path)
        elif ext == ".docx":
            content = await self._read_docx_file(file_path)
        elif ext == ".json":
            content = await self._read_text_file(file_path)
        elif ext == ".csv":
            content = await self._read_text_file(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

        # Create documents with metadata
        documents = [
            Document(
                page_content=content,
                metadata={
                    "source": filename,
                    "file_path": file_path
                }
            )
        ]

        # Split into chunks
        chunks = self.text_splitter.split_documents(documents)

        # Add chunk index to metadata
        for i, chunk in enumerate(chunks):
            chunk.metadata["chunk_index"] = i
            chunk.metadata["total_chunks"] = len(chunks)

        return chunks

    async def _read_text_file(self, file_path: str) -> str:
        """Read a text file."""
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

    async def _read_pdf_file(self, file_path: str) -> str:
        """Read a PDF file."""
        try:
            from pypdf import PdfReader

            reader = PdfReader(file_path)
            text_parts = []

            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)

            return "\n\n".join(text_parts)
        except ImportError:
            raise ImportError("pypdf is required for PDF processing. Install with: pip install pypdf")

    async def _read_docx_file(self, file_path: str) -> str:
        """Read a DOCX file."""
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(file_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text]
            return "\n\n".join(paragraphs)
        except ImportError:
            raise ImportError("python-docx is required for DOCX processing. Install with: pip install python-docx")

    def get_supported_extensions(self) -> List[str]:
        """Return list of supported file extensions."""
        return list(self.SUPPORTED_EXTENSIONS)
